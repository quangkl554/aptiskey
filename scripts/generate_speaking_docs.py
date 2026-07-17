from __future__ import annotations

import html as html_lib
import json
import re
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\QUANG\.gemini\antigravity\scratch\aptis-study")
DATA_DIR = WORKSPACE / ".tmp" / "speaking_data"
IMAGE_DIR = WORKSPACE / ".tmp" / "speaking_images"
OPTIMIZED_IMAGE_DIR = WORKSPACE / ".tmp" / "speaking_images_optimized"
OUTPUT_DIR = Path(r"C:\Users\QUANG\OneDrive - VLG\Desktop\aptis")
BASE_URL = "https://www.aptiskey.com"

OUTPUTS = {
    1: OUTPUT_DIR / "Speaking aptis part 1.docx",
    2: OUTPUT_DIR / "Speaking aptis part 2.docx",
    3: OUTPUT_DIR / "Speaking aptis part 3.docx",
    4: OUTPUT_DIR / "Speaking aptis part 4.docx",
}


def read_json(name: str) -> list[dict]:
    path = DATA_DIR / f"{name}.json"
    if not path.exists():
        raise FileNotFoundError(path)
    records = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(records, list) or not records:
        raise ValueError(f"No records in {path}")
    return records


def clean_text(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<li\b[^>]*>", "\n• ", text, flags=re.I)
    text = re.sub(r"</(?:li|p|ul|ol|div|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = html_lib.unescape(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    compact: list[str] = []
    for line in lines:
        if not line:
            if compact and compact[-1]:
                compact.append("")
        else:
            compact.append(line)
    while compact and not compact[0]:
        compact.pop(0)
    while compact and not compact[-1]:
        compact.pop()
    return "\n".join(compact)


def normalized_file_name(url_path: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", url_path.strip("/"))


def download_one(url_path: str) -> tuple[str, Path]:
    target = IMAGE_DIR / normalized_file_name(url_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    if not target.exists() or target.stat().st_size < 100:
        request = urllib.request.Request(
            BASE_URL + url_path,
            headers={"User-Agent": "Mozilla/5.0"},
        )
        with urllib.request.urlopen(request, timeout=30) as response:
            data = response.read()
        target.write_bytes(data)
    with Image.open(target) as image:
        image.verify()

    optimized = OPTIMIZED_IMAGE_DIR / (target.stem + ".jpg")
    optimized.parent.mkdir(parents=True, exist_ok=True)
    if not optimized.exists() or optimized.stat().st_mtime < target.stat().st_mtime:
        with Image.open(target) as image:
            image.load()
            if image.mode in ("RGBA", "LA"):
                rgba = image.convert("RGBA")
                background = Image.new("RGBA", rgba.size, "white")
                image = Image.alpha_composite(background, rgba).convert("RGB")
            else:
                image = image.convert("RGB")
            image.save(optimized, "JPEG", quality=88, optimize=True, progressive=True)
    with Image.open(optimized) as image:
        image.verify()
    return url_path, optimized


def download_images(part2: list[dict], part3: list[dict]) -> dict[str, Path]:
    urls = {record["urlpic1"] for record in part2}
    urls.update(record["urlpic1"] for record in part3)
    urls.update(record["urlpic2"] for record in part3)
    mapping: dict[str, Path] = {}
    with ThreadPoolExecutor(max_workers=8) as executor:
        futures = {executor.submit(download_one, url): url for url in sorted(urls)}
        for future in as_completed(futures):
            url, target = future.result()
            mapping[url] = target
    if len(mapping) != len(urls):
        raise AssertionError(f"Downloaded {len(mapping)} of {len(urls)} Speaking images")
    return mapping


def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Aptos Display"
    heading1.font.size = Pt(17)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(31, 78, 121)
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(7)
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Aptos Display"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(46, 116, 181)
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.keep_with_next = True


def new_document(part: int, subtitle: str) -> Document:
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"APTIS SPEAKING – PART {part}")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.bold = True
    run.font.size = Pt(12)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.space_after = Pt(15)
    run = source.add_run("Nguồn: AptisKey – Trang chủ → Speaking → Học theo câu hỏi (17/07/2026)")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)
    return doc


def shade_paragraph(paragraph, fill: str = "EAF3F8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_text_with_breaks(paragraph, text: str) -> None:
    pieces = clean_text(text).split("\n")
    if not pieces:
        return
    paragraph.add_run(pieces[0])
    for piece in pieces[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(piece)


def add_qa(
    doc: Document,
    question_label: str,
    question: str,
    answer_label: str,
    answer: str,
    compact: bool = False,
) -> None:
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Cm(0.25)
    q.paragraph_format.space_after = Pt(3)
    q.paragraph_format.keep_with_next = True
    run = q.add_run(question_label)
    run.bold = True
    add_text_with_breaks(q, question)

    a = doc.add_paragraph()
    a.paragraph_format.left_indent = Cm(0.4)
    a.paragraph_format.right_indent = Cm(0.15)
    a.paragraph_format.space_after = Pt(5 if compact else 7)
    shade_paragraph(a)
    run = a.add_run(answer_label)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    add_text_with_breaks(a, answer)
    if compact:
        for answer_run in a.runs:
            answer_run.font.size = Pt(10.5)


def add_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4C6E7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_picture_scaled(paragraph, path: Path, max_width_cm: float, max_height_cm: float) -> None:
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))


def build_part1(practice: list[dict], total: list[dict]) -> Document:
    doc = new_document(1, "Ngân hàng luyện tập và tổng hợp – 56 cách hỏi")
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    note.add_run(
        "AptisKey cung cấp hai danh sách Part 1 với cách diễn đạt và đáp án khác nhau; "
        "file này giữ cả hai danh sách để học offline đầy đủ."
    ).italic = True
    for section_title, rows in (("A. Luyện tập hiện tại", practice), ("B. Danh sách tổng hợp", total)):
        doc.add_paragraph(section_title, style="Heading 1")
        for index, record in enumerate(rows, start=1):
            heading = doc.add_paragraph(f"Câu {index:02d}", style="Heading 2")
            add_qa(doc, "Câu hỏi: ", record["question"], "Đáp án mẫu 1: ", record["answer1"])
            add_qa(doc, "", "", "Đáp án mẫu 2: ", record["answer2"])
            if index < len(rows):
                add_separator(doc)
    return doc


def build_part2(rows: list[dict], images: dict[str, Path]) -> Document:
    doc = new_document(2, f"{len(rows)} bộ câu hỏi có hình và đáp án mẫu")
    for index, record in enumerate(rows, start=1):
        heading = doc.add_paragraph(f"Bộ câu hỏi {index:02d}", style="Heading 1")
        heading.paragraph_format.page_break_before = index > 1
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_scaled(picture, images[record["urlpic1"]], 14.5, 8.5)
        for number in range(1, 4):
            add_qa(
                doc,
                f"Câu {number}: ",
                record[f"question{number}"],
                "Đáp án mẫu: ",
                record[f"question{number}_answer"],
            )
    return doc


def build_part3(rows: list[dict], images: dict[str, Path]) -> Document:
    doc = new_document(3, f"{len(rows)} bộ so sánh hai hình và đáp án mẫu")
    for index, record in enumerate(rows, start=1):
        heading = doc.add_paragraph(f"Bộ câu hỏi {index:02d}", style="Heading 1")
        heading.paragraph_format.page_break_before = index > 1
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        for cell, url in zip(table.rows[0].cells, (record["urlpic1"], record["urlpic2"])):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_picture_scaled(paragraph, images[url], 7.0, 7.0)
        for number in range(1, 4):
            add_qa(
                doc,
                f"Câu {number}: ",
                record[f"question{number}"],
                "Đáp án mẫu: ",
                record[f"question{number}_answer"],
                compact=True,
            )
    return doc


def build_part4(rows: list[dict]) -> Document:
    doc = new_document(4, f"{len(rows)} chủ đề kể trải nghiệm và đáp án mẫu")
    for index, record in enumerate(rows, start=1):
        title = "Mẫu form chung" if index == 1 else f"Chủ đề {index - 1:02d}"
        heading = doc.add_paragraph(title, style="Heading 1")
        heading.paragraph_format.page_break_before = index > 1
        add_qa(doc, "Câu hỏi: ", record["question"], "Đáp án mẫu: ", record["answer1"])
    return doc


def validate_doc(path: Path, part: int, expected_answers: int, expected_images: int = 0) -> None:
    reopened = Document(path)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    answer_count = sum("Đáp án mẫu" in paragraph.text for paragraph in reopened.paragraphs)
    if answer_count != expected_answers:
        raise AssertionError(f"{path.name}: {answer_count} answers != {expected_answers}")
    if len(reopened.inline_shapes) != expected_images:
        raise AssertionError(
            f"{path.name}: {len(reopened.inline_shapes)} images != {expected_images}"
        )
    if re.search(r"<(?:br|strong|span|li|ul)\b", text, flags=re.I):
        raise AssertionError(f"{path.name}: raw HTML remains")
    if f"APTIS SPEAKING – PART {part}" not in text:
        raise AssertionError(f"{path.name}: title missing")


def main() -> None:
    part1_practice = read_json("part1_practice")
    part1_total = read_json("part1_total")
    part2 = read_json("part2_practice")
    part3 = read_json("part3_practice")
    part4 = read_json("part4_practice")

    expected_counts = (28, 28, 37, 25, 53)
    actual_counts = tuple(map(len, (part1_practice, part1_total, part2, part3, part4)))
    if actual_counts != expected_counts:
        raise AssertionError(f"Speaking data counts changed: {actual_counts}")

    images = download_images(part2, part3)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    documents = {
        1: build_part1(part1_practice, part1_total),
        2: build_part2(part2, images),
        3: build_part3(part3, images),
        4: build_part4(part4),
    }
    expectations = {
        1: (112, 0),
        2: (111, 37),
        3: (75, 50),
        4: (53, 0),
    }
    for part, document in documents.items():
        path = OUTPUTS[part]
        document.save(path)
        expected_answers, expected_images = expectations[part]
        validate_doc(path, part, expected_answers, expected_images)
        print(
            f"PASS part={part} path={path} size={path.stat().st_size} "
            f"answers={expected_answers} images={expected_images}"
        )


if __name__ == "__main__":
    main()
