from __future__ import annotations

import html as html_lib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\QUANG\.gemini\antigravity\scratch\aptis-study")
DATA_DIR = WORKSPACE / ".tmp" / "writing_data"
OUTPUT_DIR = Path(r"C:\Users\QUANG\OneDrive - VLG\Desktop\aptis")

OUTPUTS = {
    1: OUTPUT_DIR / "Writting aptis part 1 - day du.docx",
    2: OUTPUT_DIR / "Writting aptis part 2.docx",
    3: OUTPUT_DIR / "Writting aptis part 3.docx",
    4: OUTPUT_DIR / "Writting aptis part 4.docx",
}


def clean_text(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = html_lib.unescape(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    compact: list[str] = []
    for line in lines:
        if line or not compact or compact[-1]:
            compact.append(line)
    return "\n".join(compact)


def load_data() -> list[dict]:
    records: list[dict] = []
    for index in range(1, 41):
        path = DATA_DIR / f"writingkey{index:03d}.json"
        if not path.exists():
            raise FileNotFoundError(f"Missing Writing dataset: {path}")
        record = json.loads(path.read_text(encoding="utf-8"))
        record["source_index"] = index
        records.append(record)
    return records


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeatable_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")

    heading = doc.styles["Heading 1"]
    heading.font.name = "Aptos Display"
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(31, 78, 121)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True


def new_document(part: int) -> Document:
    doc = Document()
    set_repeatable_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"APTIS WRITING – PART {part}")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Học câu lạc bộ – 40 chủ đề")
    run.bold = True
    run.font.size = Pt(12)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source.paragraph_format.space_after = Pt(16)
    run = source.add_run("Nguồn: AptisKey – Trang chủ → Writing → Học câu lạc bộ (17/07/2026)")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)
    return doc


def add_club_heading(doc: Document, index: int, club_name: str) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(f"Bộ đề #{index:03d} — {club_name}")


def add_labelled_paragraph(
    doc: Document,
    label: str,
    text: str,
    *,
    answer: bool = False,
    keep_with_next: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.15)
    paragraph.paragraph_format.space_after = Pt(5 if answer else 3)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if answer:
        set_cell_shading(paragraph, "EAF3F8")
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(31, 78, 121) if answer else RGBColor(0, 0, 0)
    pieces = clean_text(text).split("\n")
    if pieces:
        paragraph.add_run(pieces[0])
        for piece in pieces[1:]:
            paragraph.add_run().add_break()
            paragraph.add_run(piece)


def add_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4C6E7")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_part1(records: list[dict]) -> Document:
    doc = new_document(1)
    for position, record in enumerate(records):
        add_club_heading(doc, record["source_index"], record["club_name"])
        for number, (key, question) in enumerate(record["questions1"].items(), start=1):
            add_labelled_paragraph(doc, f"{number}. ", question, keep_with_next=True)
            add_labelled_paragraph(
                doc,
                "Đáp án mẫu: ",
                record["questions1_answer"][f"{key}_answer"],
                answer=True,
            )
        if position < len(records) - 1:
            add_separator(doc)
    return doc


def build_part2(records: list[dict]) -> Document:
    doc = new_document(2)
    for position, record in enumerate(records):
        add_club_heading(doc, record["source_index"], record["club_name"])
        question = next(iter(record["questions2"].values()))
        answer = next(iter(record["questions2_answer"].values()))
        add_labelled_paragraph(doc, "Câu hỏi: ", question, keep_with_next=True)
        add_labelled_paragraph(doc, "Đáp án mẫu: ", answer, answer=True)
        if position < len(records) - 1:
            add_separator(doc)
    return doc


def build_part3(records: list[dict]) -> Document:
    doc = new_document(3)
    for position, record in enumerate(records):
        add_club_heading(doc, record["source_index"], record["club_name"])
        for number, (key, question) in enumerate(record["questions3"].items(), start=1):
            add_labelled_paragraph(doc, f"Câu {number}: ", question, keep_with_next=True)
            add_labelled_paragraph(
                doc,
                "Đáp án mẫu: ",
                record["questions3_answer"][f"{key}_answer"],
                answer=True,
            )
        if position < len(records) - 1:
            add_separator(doc)
    return doc


def build_part4(records: list[dict]) -> Document:
    doc = new_document(4)
    for position, record in enumerate(records):
        add_club_heading(doc, record["source_index"], record["club_name"])
        add_labelled_paragraph(doc, "Tình huống chung: ", record["questions4_main"])
        add_labelled_paragraph(doc, "Email 1 – Yêu cầu: ", record["question4_1_text"], keep_with_next=True)
        add_labelled_paragraph(doc, "Đáp án mẫu Email 1: ", record["question4_1_text_answer"], answer=True)
        add_labelled_paragraph(doc, "Email 2 – Yêu cầu: ", record["question4_2_text"], keep_with_next=True)
        add_labelled_paragraph(doc, "Đáp án mẫu Email 2: ", record["question4_2_text_answer"], answer=True)
        if position < len(records) - 1:
            add_separator(doc)
    return doc


def validate_document(path: Path, part: int, records: list[dict]) -> None:
    reopened = Document(path)
    text = "\n".join(p.text for p in reopened.paragraphs)
    for index, record in enumerate(records, start=1):
        expected_heading = f"Bộ đề #{index:03d} — {record['club_name']}"
        if expected_heading not in text:
            raise AssertionError(f"Missing club heading in {path.name}: {expected_heading}")

    expected_labels = {1: 200, 2: 40, 3: 120, 4: 80}[part]
    if part == 1:
        actual_labels = sum(p.text.startswith("Đáp án mẫu: ") for p in reopened.paragraphs)
    elif part == 2:
        actual_labels = sum(p.text.startswith("Câu hỏi: ") for p in reopened.paragraphs)
    elif part == 3:
        actual_labels = sum(re.match(r"Câu [1-3]: ", p.text) is not None for p in reopened.paragraphs)
    else:
        actual_labels = sum(
            p.text.startswith("Email 1 – Yêu cầu: ") or p.text.startswith("Email 2 – Yêu cầu: ")
            for p in reopened.paragraphs
        )
    if actual_labels != expected_labels:
        raise AssertionError(
            f"Unexpected item count in {path.name}: {actual_labels} != {expected_labels}"
        )
    if "<br" in text.lower():
        raise AssertionError(f"Unconverted HTML line break in {path.name}")


def main() -> None:
    records = load_data()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    builders = {1: build_part1, 2: build_part2, 3: build_part3, 4: build_part4}
    for part, builder in builders.items():
        document = builder(records)
        path = OUTPUTS[part]
        document.save(path)
        validate_document(path, part, records)
        print(f"PASS part={part} path={path} size={path.stat().st_size}")


if __name__ == "__main__":
    main()
