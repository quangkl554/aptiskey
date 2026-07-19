from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path(r"C:\Users\QUANG\.gemini\antigravity\scratch\aptis-study")
SOURCE = WORKSPACE / "js" / "reading_db.js"
OUTPUT_DIR = Path(r"C:\Users\QUANG\OneDrive - VLG\Desktop\aptis")

ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_GREEN = "E9F5EF"
ANSWER_GREEN = "1F7A5B"
TEXT = "20252A"
MUTED = "5B6670"
WHITE = "FFFFFF"
GRID = "CAD5E0"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 0


def load_database() -> dict:
    raw = SOURCE.read_text(encoding="utf-8")
    match = re.search(r"const\s+readingDB\s*=\s*(\{.*\})\s*;\s*$", raw, re.S)
    if not match:
        raise RuntimeError(f"Cannot parse Reading database: {SOURCE}")
    payload = match.group(1)
    # The file is JSON-compatible apart from four unquoted top-level JS keys.
    payload = re.sub(r'(?m)^(\s*)(meta|part1|part2_3|part4|part5):', r'\1"\2":', payload)
    payload = re.sub(r'(?m)^(\s*)(part1Version):', r'\1"\2":', payload)
    database = json.loads(payload)

    return database


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 16, ACCENT, 18, 10),
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    add_footer(section)


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("APTIS READING · HỌC OFFLINE   |   Trang ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_editorial_header(doc: Document, title: str, subtitle: str, count_label: str) -> None:
    # First-page header pattern: editorial_cover.
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run("APTISKEY · READING · HỌC THEO CÂU HỎI")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    title_p = doc.add_paragraph(style="Title")
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(5)
    title_p.add_run(title)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(4)
    run = subtitle_p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta.add_run(count_label + "  ·  ").bold = True
    meta.add_run("Nguồn bản ghi: AptisKey  ·  Cập nhật tài liệu: 18/07/2026")
    for run in meta.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    note = doc.add_table(rows=1, cols=1)
    set_table_width(note)
    set_table_borders(note, color="B8C8D8", size=5)
    cell = note.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Cách dùng: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    r2 = p.add_run("che phần đáp án màu xanh, tự làm trước rồi mở ra đối chiếu.")
    r2.font.color.rgb = RGBColor.from_string(TEXT)


def add_text_with_breaks(paragraph, text: str, **run_kwargs) -> None:
    chunks = text.split("\n")
    for index, chunk in enumerate(chunks):
        if index:
            paragraph.add_run().add_break()
        if chunk:
            run = paragraph.add_run(chunk)
            for key, value in run_kwargs.items():
                setattr(run, key, value)


def add_template_text(paragraph, text: str, previous_text: str = "") -> None:
    if text and previous_text and previous_text[-1].isalnum() and text[0].isalnum():
        paragraph.add_run(" ")
    add_text_with_breaks(paragraph, text)


def create_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    source_num = next(
        node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_id = int(source_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_answer_badge(paragraph, answer: str, prefix="Đáp án: ") -> None:
    label = paragraph.add_run(prefix)
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(MUTED)
    answer_run = paragraph.add_run(answer)
    answer_run.bold = True
    answer_run.font.color.rgb = RGBColor.from_string(ANSWER_GREEN)


def add_standard_heading(doc: Document, number: int, topic: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(f"{number:02d}. {topic}")


def make_document(title: str, subtitle: str, count_label: str) -> Document:
    doc = Document()
    configure_styles(doc)
    core = doc.core_properties
    core.title = title
    core.subject = "Aptis Reading offline study guide"
    core.author = "Tổng hợp từ bản ghi AptisKey"
    core.keywords = "Aptis, Reading, offline, study"
    add_editorial_header(doc, title, subtitle, count_label)
    return doc


def add_part1(doc: Document, items: list[dict]) -> None:
    for item in items:
        add_standard_heading(doc, item["id"], "Sentence comprehension")

        sender = doc.add_paragraph()
        sender.paragraph_format.space_after = Pt(4)
        sender.paragraph_format.keep_with_next = True
        r = sender.add_run(item["sender"])
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

        passage = doc.add_paragraph()
        passage.paragraph_format.space_after = Pt(6)
        templates = item["template"]
        gaps = item["gaps"]
        for index, gap in enumerate(gaps):
            previous_text = passage.text
            add_template_text(passage, templates[index], previous_text)
            if passage.text and passage.text[-1].isalnum():
                passage.add_run(" ")
            answer_run = passage.add_run(gap["correct"])
            answer_run.bold = True
            answer_run.underline = True
            answer_run.font.color.rgb = RGBColor.from_string(ANSWER_GREEN)
        add_template_text(passage, templates[-1], gap["correct"])

        table = doc.add_table(rows=1, cols=3)
        set_table_width(table)
        set_table_borders(table)
        headers = ("Chỗ trống", "Các lựa chọn", "Đáp án đúng")
        widths = (950, 6200, 2210)
        header_row = table.rows[0]
        set_repeat_table_header(header_row)
        prevent_row_split(header_row)
        for idx, (cell, header, width) in enumerate(zip(header_row.cells, headers, widths)):
            set_cell_width(cell, width)
            set_cell_shading(cell, LIGHT_BLUE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(header)
            rr.bold = True
            rr.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

        for index, gap in enumerate(gaps, 1):
            row = table.add_row()
            prevent_row_split(row)
            values = (str(index), " / ".join(gap["options"]), gap["correct"])
            for col_idx, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
                set_cell_width(cell, width)
                set_cell_margins(cell)
                if col_idx == 2:
                    set_cell_shading(cell, PALE_GREEN)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                rr = p.add_run(value)
                if col_idx == 2:
                    rr.bold = True
                    rr.font.color.rgb = RGBColor.from_string(ANSWER_GREEN)

        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_part2_3(doc: Document, items: list[dict]) -> None:
    for index, item in enumerate(items, 1):
        add_standard_heading(doc, index, item["topic"])
        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(4)
        intro.add_run("Thứ tự đúng:").bold = True
        num_id = create_numbering_instance(doc)
        for order, sentence in enumerate(item["correctOrder"], 1):
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, num_id)
            p.paragraph_format.keep_together = True
            run = p.add_run(sentence)
            if order == 1:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)


def add_part4(doc: Document, items: list[dict]) -> None:
    for index, item in enumerate(items, 1):
        add_standard_heading(doc, index, item["topic"])

        for letter in ("A", "B", "C", "D"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.keep_together = True
            badge = p.add_run(f"{letter}  ")
            badge.bold = True
            badge.font.color.rgb = RGBColor.from_string(ACCENT)
            p.add_run(item["texts"][letter])

        q_heading = doc.add_paragraph(style="Heading 3")
        q_heading.add_run("Câu hỏi và đáp án")

        table = doc.add_table(rows=1, cols=3)
        set_table_width(table)
        set_table_borders(table)
        widths = (650, 7550, 1160)
        headers = ("#", "Câu hỏi", "Đáp án")
        header_row = table.rows[0]
        set_repeat_table_header(header_row)
        prevent_row_split(header_row)
        for col_idx, (cell, header, width) in enumerate(zip(header_row.cells, headers, widths)):
            set_cell_width(cell, width)
            set_cell_shading(cell, LIGHT_BLUE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            rr = p.add_run(header)
            rr.bold = True
            rr.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

        for q_number, statement in enumerate(item["statements"], 1):
            row = table.add_row()
            prevent_row_split(row)
            values = (str(q_number), statement["statement"], statement["correct"])
            for col_idx, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
                set_cell_width(cell, width)
                set_cell_margins(cell)
                if col_idx == 2:
                    set_cell_shading(cell, PALE_GREEN)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                rr = p.add_run(value)
                if col_idx == 2:
                    rr.bold = True
                    rr.font.color.rgb = RGBColor.from_string(ANSWER_GREEN)


def add_part5(doc: Document, items: list[dict]) -> None:
    for index, item in enumerate(items, 1):
        add_standard_heading(doc, index, item["topic"])

        bank = doc.add_table(rows=1, cols=1)
        set_table_width(bank)
        set_table_borders(bank, color="B8C8D8", size=5)
        cell = bank.cell(0, 0)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Ngân hàng tiêu đề: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
        p.add_run(" · ".join(item["headings"]))

        for paragraph in item["paragraphs"]:
            h = doc.add_paragraph(style="Heading 3")
            h.add_run(f"Đoạn {paragraph['num']}  ·  ")
            ans = h.add_run(paragraph["correct"])
            ans.font.color.rgb = RGBColor.from_string(ANSWER_GREEN)

            p = doc.add_paragraph(paragraph["text"])
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.widow_control = True


def validate_database(db: dict) -> dict:
    expected_counts = {"part1": 45, "part2_3": 39, "part4": 14, "part5": 11}
    actual_counts = {key: len(db.get(key, [])) for key in expected_counts}
    if actual_counts != expected_counts:
        raise AssertionError(f"Unexpected Reading counts: {actual_counts}")

    expected_ids = list(range(1, expected_counts["part1"] + 1))
    actual_ids = [item.get("id") for item in db["part1"]]
    if actual_ids != expected_ids:
        raise AssertionError(f"Part 1 ids must be sequential 1-45; found {actual_ids}")

    seen_part1_sets = set()
    for item in db["part1"]:
        fingerprint = json.dumps(
            {key: value for key, value in item.items() if key != "id"},
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        if fingerprint in seen_part1_sets:
            raise AssertionError(f"Duplicate Part 1 question set at id {item.get('id')}")
        seen_part1_sets.add(fingerprint)
        if not item.get("sender"):
            raise AssertionError(f"Part 1 question {item.get('id')} has no sender")
        if len(item.get("gaps", [])) != 5 or len(item.get("template", [])) != 6:
            raise AssertionError(
                f"Part 1 question {item.get('id')} must have 5 gaps and 6 template segments"
            )
        for gap_index, gap in enumerate(item["gaps"], 1):
            options = gap.get("options", [])
            if len(options) != 3 or len(set(options)) != 3:
                raise AssertionError(
                    f"Part 1 question {item['id']}, gap {gap_index} must have "
                    f"exactly 3 unique options"
                )
            if options.count(gap.get("correct")) != 1:
                raise AssertionError(
                    f"Part 1 question {item['id']}, gap {gap_index} has an invalid correct answer"
                )

    for item in db["part2_3"]:
        assert item["topic"]
        assert len(item["correctOrder"]) == 5
        assert len(set(item["correctOrder"])) == 5

    for item in db["part4"]:
        assert item["topic"]
        assert set(item["texts"]) == {"A", "B", "C", "D"}
        assert all(item["texts"].values())
        assert len(item["statements"]) == 7
        assert all(x["correct"] in {"A", "B", "C", "D"} for x in item["statements"])

    for item in db["part5"]:
        assert item["topic"]
        assert len(item["paragraphs"]) == 7
        assert len(item["headings"]) == 7
        assert all(x["correct"] in item["headings"] for x in item["paragraphs"])
        assert all(x["text"] for x in item["paragraphs"])

    return actual_counts


def save_and_reopen(doc: Document, path: Path, part_key: str, expected_items: int) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    check = Document(path)
    if not check.paragraphs:
        raise AssertionError(f"Generated document is empty: {path}")
    if part_key == "part1":
        heading_count = sum(
            1
            for paragraph in check.paragraphs
            if paragraph.style and paragraph.style.name == "Heading 2"
        )
        question_tables = [
            table
            for table in check.tables
            if len(table.columns) == 3
            and len(table.rows) >= 1
            and [cell.text.strip() for cell in table.rows[0].cells]
            == ["Chỗ trống", "Các lựa chọn", "Đáp án đúng"]
        ]
        answer_rows = sum(len(table.rows) - 1 for table in question_tables)
        expected_answer_rows = expected_items * 5
        if heading_count != expected_items or len(question_tables) != expected_items:
            raise AssertionError(
                f"Part 1 DOCX structure mismatch: headings={heading_count}, "
                f"tables={len(question_tables)}, expected={expected_items}"
            )
        if answer_rows != expected_answer_rows:
            raise AssertionError(
                f"Part 1 DOCX answer rows={answer_rows}, expected={expected_answer_rows}"
            )


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Reading DOCX study banks")
    parser.add_argument("--part", choices=("part1", "all"), default="part1")
    args = parser.parse_args()

    db = load_database()
    counts = validate_database(db)

    specs = [
        (
            "part1",
            "Reading aptis part 1.docx",
            "APTIS READING – PART 1",
            "Sentence comprehension · Chọn từ phù hợp với chỗ trống",
            (
                f"{counts['part1']} bộ câu hỏi · "
                f"{counts['part1'] * 5} chỗ trống"
            ),
            add_part1,
        ),
        (
            "part2_3",
            "Reading aptis part 2&3.docx",
            "APTIS READING – PART 2 & 3",
            "Text cohesion · Sắp xếp câu theo đúng trình tự",
            f"{counts['part2_3']} chủ đề · {counts['part2_3'] * 5} câu",
            add_part2_3,
        ),
        (
            "part4",
            "Reading aptis part 4.docx",
            "APTIS READING – PART 4",
            "Opinion matching · Ghép quan điểm với người nói",
            f"{counts['part4']} chủ đề · {counts['part4'] * 7} câu hỏi",
            add_part4,
        ),
        (
            "part5",
            "Reading aptis part 5.docx",
            "APTIS READING – PART 5",
            "Long text comprehension · Ghép tiêu đề với đoạn văn",
            f"{counts['part5']} chủ đề · {counts['part5'] * 7} đoạn văn",
            add_part5,
        ),
    ]
    if args.part == "part1":
        specs = [specs[0]]

    outputs = []
    for key, filename, title, subtitle, count_label, builder in specs:
        doc = make_document(title, subtitle, count_label)
        builder(doc, db[key])
        path = OUTPUT_DIR / filename
        save_and_reopen(doc, path, key, counts[key])
        outputs.append({"path": str(path), "bytes": path.stat().st_size, "items": counts[key]})

    print(json.dumps({"source": str(SOURCE), "counts": counts, "outputs": outputs}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
